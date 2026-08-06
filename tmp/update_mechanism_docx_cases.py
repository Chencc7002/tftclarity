from copy import deepcopy
from pathlib import Path

from docx import Document


DOCX_PATH = Path("docs/mechanism-classification/机制分类定义与案例-整理版.docx")


def set_case_table(table, values):
    for row in table.rows:
        if not row.cells:
            continue
        key = row.cells[0].text.strip()
        if key in values and len(row.cells) > 1:
            row.cells[1].text = values[key]


def add_negative_row(table, case_text, judgment, reason):
    template = table.rows[-1]
    row = table.add_row()
    for index, cell in enumerate(row.cells):
        if index < len(template.cells) and template.cells[index]._tc.tcPr is not None:
            cell._tc.remove(cell._tc.tcPr)
            cell._tc.insert(0, deepcopy(template.cells[index]._tc.tcPr))
    row.cells[0].text = case_text
    row.cells[1].text = judgment
    row.cells[2].text = reason


document = Document(DOCX_PATH)

for paragraph in document.paragraphs:
    if paragraph.text.strip() in {"D-011  命运祭司", "D-011  命运祭祀"}:
        paragraph.text = "D-011  命运祭司（命运祭祀）"

set_case_table(document.tables[11], {
    "分类": "发育",
    "实体类型": "羁绊",
    "触发条件": "备战席克隆栏位完成棋子克隆",
    "累计/升级": "克隆进度完成后获得目标棋子的 1 星版本和金币",
    "结论": "完成克隆后直接获得棋子和金币，属于资源产出。",
    "判定理由": "棋子和金币均属于发育资源；克隆进度服务于资源产出，不单独算作永久战力成长。",
    "持久性": "not_applicable",
    "置信度": "0.99",
})

set_case_table(document.tables[21], {
    "分类": "发育",
    "实体类型": "羁绊",
    "触发条件": "输掉玩家对战、成员参与击杀并累计研究点",
    "累计/升级": "研究点达到门槛后制造武器原型",
    "结论": "研究点达到门槛后制造并提供装备，属于资源与战利品产出。",
    "判定理由": "研究点和武器原型属于发育资源；获得装备本身不计为实体永久战力成长。",
    "持久性": "not_applicable",
    "置信度": "0.99",
})

negative_table = document.tables[24]
existing_cases = {row.cells[0].text for row in negative_table.rows if row.cells}
if not any("观星者:泉水" in text for text in existing_cases):
    add_negative_row(
        negative_table,
        "N-011\n观星者:泉水",
        "成长：否\n发育：否",
        "物理与法术加成只在当前战斗内每 2 秒叠加，战斗结束后重置，不属于跨回合永久成长。",
    )
if not any("太空律动" in text for text in existing_cases):
    add_negative_row(
        negative_table,
        "N-012\n太空律动",
        "成长：否\n发育：否",
        "战斗属性只在太空律动状态下逐秒叠加，不会带入后续回合，因此不属于成长。",
    )

document.save(DOCX_PATH)
